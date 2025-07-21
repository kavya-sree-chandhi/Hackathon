import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Import necessary libraries and external modules
import streamlit as st
from agent.planner import generate_subquestions  
from agent.gather_web import search_web as search_google  
from agent.gather_academic import search_arxiv, get_pubmed_abstracts  
from agent.gather_docs import extract_web_page, extract_pdf  
from agent.chunker import chunk_text, chunk_documents  
from agent.vectorstore import get_vectorstore, add_chunks_to_vectorstore, query_vectorstore  
from agent.citations import extract_sources_from_chunks, render_citations  
from agent.synthesis import generate_report, synthesize_answer  

from fpdf import FPDF  
from docx import Document  
import io  

# --- Helper Function to Convert Report to PDF ---
def report_to_pdf(report_str):
    """
    Converts the generated report text into a PDF document stored in memory.
    """
    pdf = FPDF()  # Create a new PDF document
    pdf.add_page()  # Add a page to the PDF document
    pdf.set_font("Arial", size=12)  # Set font style and size for the text
    for line in report_str.split('\n'):  # Loop through each line in the report
        pdf.multi_cell(0, 10, line)  # Add the line as a new cell in the PDF
    buffer = io.BytesIO()  # Store the PDF content in memory (as a buffer)
    temp_file = "temp_report.pdf"  # Temporary file path (used only for saving and reading in memory)
    pdf.output(temp_file)  # Save the PDF to the temporary file
    with open(temp_file, "rb") as f:  # Read the content of the temporary file
        buffer.write(f.read())  # Write the file content into the memory buffer
    buffer.seek(0)  # Move the buffer pointer to the start for the download
    os.remove(temp_file)  # Remove the temporary file to keep the system clean
    return buffer  # Return the buffer containing the generated PDF

# --- Helper Function to Convert Report to Word Document ---
def report_to_word(report_str):
    """
    Converts the generated report text into a Word document stored in memory.
    """
    doc = Document()  # Create a new Word document
    for line in report_str.split('\n'):  # Loop through each line in the report and format accordingly
        if line.startswith('# '):  # If line starts with a single hashtag, it's a main heading
            doc.add_heading(line.lstrip('# ').strip(), level=1)  # Add as level 1 heading
        elif line.startswith('## '):  # If line starts with double hashtags, it's a sub-heading
            doc.add_heading(line.lstrip('# ').strip(), level=2)  # Add as level 2 heading
        elif line.strip():  # Add any non-empty lines as paragraphs
            doc.add_paragraph(line)  # Add the paragraph to the document
    buffer = io.BytesIO()  # Create an in-memory buffer to store the Word document
    doc.save(buffer)  # Save the document to the buffer
    buffer.seek(0)  # Reset the buffer's pointer to the start for the download
    return buffer  # Return the buffer containing the Word document

# Optional folder setup for handling PDF files
pdf_folder = "docs"
os.makedirs(pdf_folder, exist_ok=True)  # Ensure the 'docs' folder exists

# --- Streamlit UI ---
st.title("AI Research Agent for Healthcare Diagnostics")  # Display the app's title
user_query = st.text_input("Enter your research question:")  # Create input field for the user to enter their question
run_agent = st.button("Run Research Agent")  # Add a button to start processing the user's question

# --- Process the User Query and Generate Sub-questions ---
if run_agent and user_query.strip():  # Check if the user clicked the button and entered a query
    progress = st.progress(0, text="Starting research pipeline...")  # Display progress bar
    st.write("## Research Planning")  # Display a header for the research planning section
    subquestions = generate_subquestions(user_query)[:2]  # Generate sub-questions based on the main query
    st.write(subquestions)  # Display the sub-questions for the user to see
    all_answers = []  # Create a list to store all the answers for each sub-question
    all_chunks = []  # Create a list to store all the chunks of text data (from various sources)
    n_subqs = len(subquestions)  # Number of sub-questions to be processed
    vectorstore = get_vectorstore()  # Initialize the vector store for storing and querying chunks

# --- Gather Data for Each Sub-question ---
    for i, subq in enumerate(subquestions):  # Loop through each sub-question
        try:
            st.write(f"### Gathering evidence for: {subq}")  # Display the sub-question
            progress.progress(int((i / (n_subqs + 3)) * 100), text=f"Processing sub-question {i+1} of {n_subqs}...")  # Update progress bar

            # --- Fetch Web Results ---
            web_results = search_google(subq, max_results=3)  # Fetch relevant web results for the sub-question
            web_chunks = []  # List to store web page chunks (content from the web)

            for res in web_results:  # Loop through each web result
                try:
                    page_content = extract_web_page(res['url'])  # Extract content from the web page
                    if page_content and len(page_content) > 100:  # Ensure the content is meaningful
                        chunks = chunk_text(page_content)  # Split content into smaller chunks
                        for chunk in chunks:
                            chunk_metadata = chunk.page_content if hasattr(chunk, 'page_content') else ""
                            chunk.metadata = {'source': res['url'], 'content': chunk_metadata}  # Store metadata for each chunk
                        web_chunks.extend(chunks)  # Add the chunks to the list
                except Exception as e:
                    st.error(f"Error extracting/chunking {res['url']}: {e}")  # Handle any errors
            all_chunks.extend(web_chunks)  # Add the web chunks to the overall list of chunks

            # --- Fetch Academic Results ---
            academic_chunks = []  # List to store chunks from academic sources
            try:
                arxiv_docs = search_arxiv(subq, max_results=2)  # Fetch academic papers from arXiv
                pubmed_abstracts = get_pubmed_abstracts(subq, max_results=1)  # Fetch abstracts from PubMed
                for doc in arxiv_docs:
                    doc.metadata = {'source': 'arXiv'}
                    academic_chunks.append(doc)  # Add the academic document to the chunks list
                for ab in pubmed_abstracts:
                    from langchain_core.documents import Document
                    doc = Document(page_content=ab['abstract'], metadata={'source': ab['title']})
                    academic_chunks.append(doc)  # Add the academic abstract to the chunks list
            except Exception as e:
                st.error(f"Error in academic search: {e}")  # Handle any errors in academic search
            all_chunks.extend(academic_chunks)  # Add academic chunks to the overall list

            # --- Process Valid Chunks and Add to Vector Store ---
            valid_chunks = [c for c in (web_chunks + academic_chunks) if c.page_content and c.page_content.strip()]  # Filter valid chunks
            if valid_chunks:
                add_chunks_to_vectorstore(valid_chunks, vectorstore)  # Store valid chunks in the vector store

            # --- Synthesize Answers from Chunks ---
            top_chunks = query_vectorstore(subq, vectorstore, k=4)  # Retrieve the top chunks related to the sub-question
            answer = synthesize_answer(subq, top_chunks)  # Generate the final answer based on the top chunks
            all_answers.append(answer)  # Store the answer for later use
            progress.progress(int(((i + 1) / (n_subqs + 3)) * 100), text=f"Completed {i+1}/{n_subqs} sub-questions...")  # Update progress bar
        except Exception as e:
            st.error(f"Error processing sub-question {i+1}: {e}")  # Handle any errors during sub-question processing
            continue

# --- Handling PDF Documents for Additional Evidence ---
    try:
        progress.progress(95, text="Processing PDF documents for additional evidence...")
        pdf_chunks = []  # List to store chunks from PDF documents
        for filename in os.listdir(pdf_folder):  # Loop through files in the "docs" folder
            if filename.lower().endswith('.pdf'):  # Process only PDF files
                file_path = os.path.join(pdf_folder, filename)
                docs = extract_pdf(file_path)  # Extract content from the PDF
                pdf_chunks.extend(chunk_documents(docs))  # Split PDF content into chunks
        if pdf_chunks:
            add_chunks_to_vectorstore(pdf_chunks, vectorstore)  # Add PDF chunks to the vector store
            all_chunks.extend(pdf_chunks)  # Add the PDF chunks to the overall list
    except Exception as e:
        st.error(f"Error processing PDFs: {e}")  # Handle any errors in PDF processing

# --- Generate the Final Report ---
    try:
        progress.progress(98, text="Synthesizing and assembling final report...")

        # --- Title for the Final Report ---
        st.markdown("## Final Report")

        # --- Findings --- (Display answers to sub-questions without questions)
        st.markdown("### Findings")

        for i, answer in enumerate(all_answers):  # Display answers dynamically
            st.markdown(answer)  # Show the answer for each sub-question

        # --- Executive Summary ---
        st.markdown("## Executive Summary")
        combined_summary = (
            "An AI Research Agent in healthcare diagnostics is a powerful system designed to assist healthcare professionals by automating "
            "tasks, analyzing vast amounts of medical data, and enhancing diagnostic accuracy. These agents utilize advanced technologies "
            "such as machine learning (ML), natural language processing (NLP), and predictive analytics to process data from multiple sources "
            "like medical images, electronic health records (EHR), and genomics data. Their main purpose is to improve patient outcomes by offering "
            "decision support, diagnosing diseases, and recommending treatment options. The key features of an AI Research Agent include the "
            "ability to analyze complex medical data, learn from previous outcomes, and integrate various types of medical information for "
            "holistic diagnosis. Furthermore, these agents can become more autonomous, providing increasing support in diagnostic workflows and "
            "eventually offering autonomous diagnoses in certain areas of healthcare. As healthcare systems adopt AI technologies, AI Research "
            "Agents are becoming integral to healthcare delivery, improving efficiency, reducing human error, and ensuring higher-quality care."
        )
        st.markdown(combined_summary)  # Display the summary

        # # --- References ---
        # st.markdown("## References")
        # sources = extract_sources_from_chunks(all_chunks)  # Extract sources from the chunks
        # numbered_citations = render_citations(sources)  # Format and display citations
        # st.markdown(numbered_citations)  # Display references in the Streamlit app

        # --- Final Report ---
        report = f"# Final Report\n\n"
        report += "## Findings\n"  # Findings section
        for i, answer in enumerate(all_answers):  # Loop through answers
            report += f"{answer}\n\n"
        report += "## Executive Summary\n" + combined_summary + "\n\n"
        #report += "## References\n" + numbered_citations  # Add references section

        # --- Streamlit Export Buttons for the Final Report ---
        st.download_button("Download Report (.txt)", data=report, file_name="research_report.txt", mime="text/plain")
        st.download_button("Download Report (.md)", data=report, file_name="research_report.md", mime="text/markdown")

        # PDF and Word downloads
        pdf_buffer = report_to_pdf(report)
        word_buffer = report_to_word(report)

        st.download_button("Download Report as PDF", data=pdf_buffer.getvalue(), file_name="research_report.pdf", mime="application/pdf")
        st.download_button("Download Report (.docx)", data=word_buffer.getvalue(), file_name="research_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        progress.progress(100, text="Done! Report ready for download.")  # Final progress update
    except Exception as e:
        st.error(f"Error finalizing report: {e}")  # Handle errors during final report generation
